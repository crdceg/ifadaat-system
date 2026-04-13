import sys
import os
import pandas as pd
from docx import Document
from datetime import datetime

from openpyxl import load_workbook
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.utils import get_column_letter

from PySide6.QtWidgets import (
    QApplication, QWidget, QLabel, QLineEdit, QTextEdit,
    QPushButton, QVBoxLayout, QHBoxLayout, QComboBox,
    QMessageBox, QTableWidget, QTableWidgetItem
)
from PySide6.QtCore import Qt

# ====== إعدادات ======

OUTPUT_DIR = "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

SERVICE_FILES = {
    "مراجعة لغوية": os.path.join("databases", "language.xlsx"),
    "مراجعة إحصائية": os.path.join("databases", "statistics.xlsx"),
}

TEMPLATES = {
    "مراجعة لغوية": "إفادة لغوية.docx",
    "مراجعة إحصائية": "إفادة إحصائية.docx",
}

PREFIX_MAP = {
    "مراجعة لغوية": "L",
    "مراجعة إحصائية": "S",
}

DEGREES = ["الماجستير الأكاديمي", "الدكتوراه الأكاديمي", "الماجستير المهني", "الدكتوراه المهني"]

DEPARTMENTS = [
    "إدارة الأعمال",
    "إدارة الأعمال الدولية",
    "إدارة البنوك",
    "إدارة المنشآت البترولية والطاقة",
    "إدارة الموارد البشرية",
    "الإدارة العامة والمحلية",
    "إدارة الإنتاج",
    "الاقتصاد",
    "المحاسبة",
    "التأمين والإحصاء",
    "نظم المعلومات الإدارية",
    "التمويل والاستثمار",
    "التسويق والتجارة الإلكترونية",
]

# ====== Excel ======

def format_excel(file_path):
    wb = load_workbook(file_path)
    ws = wb.active

    last_row = ws.max_row
    last_col = ws.max_column
    table_ref = f"A1:{get_column_letter(last_col)}{last_row}"

    for t in list(ws.tables):
        del ws.tables[t]

    table = Table(displayName="DataTable", ref=table_ref)
    style = TableStyleInfo(name="TableStyleMedium2", showRowStripes=True)
    table.tableStyleInfo = style
    ws.add_table(table)

    ws.freeze_panes = "A2"

    for col in ws.columns:
        max_length = max((len(str(c.value)) for c in col if c.value), default=0)
        ws.column_dimensions[col[0].column_letter].width = max_length + 3

    wb.save(file_path)

# ====== Core ======

def generate_code(service):
    prefix = PREFIX_MAP[service]
    year = datetime.now().strftime("%y")
    file_path = SERVICE_FILES[service]

    if not os.path.exists(file_path):
        return f"{prefix}-{year}001"

    df = pd.read_excel(file_path, dtype=str)

    max_number = 0
    for code in df["CODE"]:
        try:
            number = int(code.split("-")[1][2:])
            max_number = max(max_number, number)
        except:
            continue

    return f"{prefix}-{year}{str(max_number + 1).zfill(3)}"


def check_duplicate(service, name):
    file_path = SERVICE_FILES[service]
    if not os.path.exists(file_path):
        return False

    df = pd.read_excel(file_path, dtype=str)
    return any(df["RESEARCHER"].str.lower() == name.lower())


def replace_placeholders(doc, mapping):
    # 🔹 paragraphs
    for p in doc.paragraphs:
        for key, value in mapping.items():
            if key in p.text:
                for run in p.runs:
                    run.text = run.text.replace(key, value)

    # 🔹 tables (المهم 🔥)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in mapping.items():
                        if key in p.text:
                            for run in p.runs:
                                run.text = run.text.replace(key, value)

def save_to_excel(service, data):
    file_path = SERVICE_FILES[service]

    if os.path.exists(file_path):
        df = pd.read_excel(file_path, dtype=str)
    else:
        df = pd.DataFrame(columns=data.keys())

    df = pd.concat([df, pd.DataFrame([data])], ignore_index=True)
    df.to_excel(file_path, index=False)

    format_excel(file_path)


def generate_doc(service, data):
    doc = Document(TEMPLATES[service])
    replace_placeholders(doc, data)

    filename = f"{data['{{CODE}}']} - {data['{{RESEARCHER}}']}.docx"
    doc.save(os.path.join(OUTPUT_DIR, filename))

# ====== Edit ======

class EditForm(QWidget):
    def __init__(self, record, file_path, service):
        super().__init__()

        self.record = record
        self.file_path = file_path
        self.service = service

        self.setWindowTitle("تعديل")
        self.setLayoutDirection(Qt.RightToLeft)

        layout = QVBoxLayout()

        self.name = QLineEdit(record["RESEARCHER"])

        self.degree = QComboBox()
        self.degree.addItems(DEGREES)
        self.degree.setCurrentText(record["DEGREE"])

        self.department = QComboBox()
        self.department.addItems(DEPARTMENTS)
        self.department.setCurrentText(record["DEPARTMENT"])

        self.title = QTextEdit(record["TITLE"])

        layout.addWidget(QLabel("اسم الباحث"))
        layout.addWidget(self.name)

        layout.addWidget(QLabel("الدرجة"))
        layout.addWidget(self.degree)

        layout.addWidget(QLabel("القسم"))
        layout.addWidget(self.department)

        layout.addWidget(QLabel("العنوان"))
        layout.addWidget(self.title)

        btn = QPushButton("حفظ + إعادة إصدار")
        btn.clicked.connect(self.save)
        layout.addWidget(btn)

        self.setLayout(layout)

    def save(self):
        df = pd.read_excel(self.file_path, dtype=str)

        code = self.record["CODE"]

        df.loc[df["CODE"] == code, "RESEARCHER"] = self.name.text()
        df.loc[df["CODE"] == code, "DEGREE"] = self.degree.currentText()
        df.loc[df["CODE"] == code, "DEPARTMENT"] = self.department.currentText()
        df.loc[df["CODE"] == code, "TITLE"] = self.title.toPlainText()

        df.to_excel(self.file_path, index=False)
        format_excel(self.file_path)

        doc_data = {
            "{{CODE}}": code,
            "{{RESEARCHER}}": self.name.text(),
            "{{DEGREE}}": self.degree.currentText(),
            "{{DEPARTMENT}}": self.department.currentText(),
            "{{TITLE}}": self.title.toPlainText(),
        }

        generate_doc(self.service, doc_data)

        QMessageBox.information(self, "تم", "تم التعديل وإعادة الإصدار")
        self.close()

# ====== Search ======

class SearchWindow(QWidget):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("بحث")
        self.setLayoutDirection(Qt.RightToLeft)

        layout = QVBoxLayout()

        self.service = QComboBox()
        self.service.addItems(SERVICE_FILES.keys())

        self.search_input = QLineEdit()

        btn = QPushButton("بحث")
        btn.clicked.connect(self.search)

        self.table = QTableWidget()

        edit_btn = QPushButton("تعديل")
        edit_btn.clicked.connect(self.edit)

        delete_btn = QPushButton("حذف")
        delete_btn.clicked.connect(self.delete)

        layout.addWidget(self.service)
        layout.addWidget(self.search_input)
        layout.addWidget(btn)
        layout.addWidget(self.table)
        layout.addWidget(edit_btn)
        layout.addWidget(delete_btn)

        self.setLayout(layout)

    def search(self):
        service = self.service.currentText()
        file_path = SERVICE_FILES[service]

        if not os.path.exists(file_path):
            QMessageBox.warning(self, "خطأ", "لا يوجد بيانات")
            return

        df = pd.read_excel(file_path, dtype=str)

        keyword = self.search_input.text().lower()
        df = df[df["RESEARCHER"].str.lower().str.contains(keyword)]

        self.data = df
        self.file_path = file_path
        self.service_name = service

        self.table.setRowCount(len(df))
        self.table.setColumnCount(len(df.columns))
        self.table.setHorizontalHeaderLabels(df.columns)

        for i in range(len(df)):
            for j in range(len(df.columns)):
                self.table.setItem(i, j, QTableWidgetItem(str(df.iloc[i, j])))

    def edit(self):
        row = self.table.currentRow()
        if row == -1:
            return

        record = self.data.iloc[row]
        self.edit_form = EditForm(record, self.file_path, self.service_name)
        self.edit_form.show()

    def delete(self):
        row = self.table.currentRow()
        if row == -1:
            return

        code = self.table.item(row, 0).text()

        df = pd.read_excel(self.file_path, dtype=str)
        df = df[df["CODE"] != code]

        df.to_excel(self.file_path, index=False)
        format_excel(self.file_path)

        QMessageBox.information(self, "تم", "تم الحذف")
        self.search()

# ====== Main App ======

class App(QWidget):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("نظام الإفادات")
        self.resize(800, 600)
        self.setMinimumSize(800, 600)
        self.setLayoutDirection(Qt.RightToLeft)

        layout = QVBoxLayout()

        # نوع الإصدار 🔥
        self.issue_type = QComboBox()
        self.issue_type.addItems([
            "مراجعة لغوية",
            "مراجعة إحصائية",
            "الاثنين معًا"
        ])

        self.name = QLineEdit()

        self.degree = QComboBox()
        self.degree.addItems(DEGREES)

        self.department = QComboBox()
        self.department.addItems(DEPARTMENTS)

        self.title = QTextEdit()

        btn = QPushButton("إصدار")
        btn.clicked.connect(self.submit)

        search_btn = QPushButton("بحث / تعديل")
        search_btn.clicked.connect(self.open_search)

        layout.addWidget(QLabel("نوع الإصدار"))
        layout.addWidget(self.issue_type)

        layout.addWidget(QLabel("اسم الباحث"))
        layout.addWidget(self.name)

        layout.addWidget(QLabel("الدرجة"))
        layout.addWidget(self.degree)

        layout.addWidget(QLabel("القسم"))
        layout.addWidget(self.department)

        layout.addWidget(QLabel("العنوان"))
        layout.addWidget(self.title)

        layout.addWidget(btn)
        layout.addWidget(search_btn)

        self.setLayout(layout)

    def submit(self):
        issue_type = self.issue_type.currentText()

        if issue_type == "الاثنين معًا":
            services = ["مراجعة لغوية", "مراجعة إحصائية"]
        else:
            services = [issue_type]

        name = self.name.text().strip()

        if not name:
            QMessageBox.warning(self, "خطأ", "ادخل اسم الباحث")
            return

        for service in services:

            if check_duplicate(service, name):
                if QMessageBox.question(self, "تنبيه", f"{service} مكرر، تكمل؟") != QMessageBox.Yes:
                    return

            code = generate_code(service)

            data = {
                "CODE": code,
                "RESEARCHER": name,
                "DEGREE": self.degree.currentText(),
                "DEPARTMENT": self.department.currentText(),
                "TITLE": self.title.toPlainText(),
                "DATE": datetime.now().strftime("%d/%m/%Y")
            }

            doc_data = {
                "{{CODE}}": code,
                "{{RESEARCHER}}": name,
                "{{DEGREE}}": self.degree.currentText(),
                "{{DEPARTMENT}}": self.department.currentText(),
                "{{TITLE}}": self.title.toPlainText(),
            }

            save_to_excel(service, data)
            generate_doc(service, doc_data)

        QMessageBox.information(self, "تم", "تم إصدار الإفادات بنجاح")

        self.name.clear()
        self.title.clear()

    def open_search(self):
        self.w = SearchWindow()
        self.w.show()

# ====== Run ======

if __name__ == "__main__":
    app = QApplication(sys.argv)
    w = App()
    w.show()
    sys.exit(app.exec())